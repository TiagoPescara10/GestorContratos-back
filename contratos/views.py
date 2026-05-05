import logging
import os
from decimal import Decimal
import io
import re
from datetime import datetime, timedelta
import requests as req

from django.shortcuts import get_object_or_404
from django.utils import timezone
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.core.files.base import ContentFile
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
import cloudinary
from cloudinary.utils import cloudinary_url

from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.filters import SearchFilter, OrderingFilter
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from utils.numero_a_letras import convertir_monto_a_letras

from .models import Contrato, EstadoMensual
from .serializers import (
    ContratoListSerializer, ContratoDetailSerializer,
    EstadoMensualSerializer, EstadoMensualUpdateSerializer,
    AplicarAumentoSerializer, ConfirmarAumentoSerializer, AplicarMoraSerializer,
    ReciboSerializer,
)
from .filters import ContratoFilter
from . import services
from indices.client import obtener_indice

logger = logging.getLogger('contratos')


class ContratoViewSet(viewsets.ModelViewSet):
    """
    CRUD completo de contratos + acciones especiales.

    list:   GET  /api/contratos/
    create: POST /api/contratos/
    retrieve: GET  /api/contratos/{id}/
    update: PUT  /api/contratos/{id}/
    destroy: DELETE /api/contratos/{id}/   (soft delete)
    """
    queryset        = Contrato.objects.filter(eliminado=False)
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_class = ContratoFilter
    search_fields   = ['inquilinoNombre', 'inquilinoDni', 'propietarioNombre', 'localidad']
    ordering_fields = ['fechaInicio', 'fechaFin', 'valorMensual', 'localidad', 'createdAt']
    ordering        = ['-createdAt']
    parser_classes  = [MultiPartParser, FormParser, JSONParser]

    def get_queryset(self):
        """
        Filter contracts by authenticated user.
        Superusers can see all contracts.
        """
        # Verificar que el usuario esté autenticado
        if not self.request.user or not self.request.user.is_authenticated:
            from rest_framework.exceptions import AuthenticationFailed
            raise AuthenticationFailed('Usuario no autenticado')
            
        if self.request.user.is_superuser:
            return Contrato.objects.filter(eliminado=False)
        return Contrato.objects.filter(usuario=self.request.user, eliminado=False)

    def get_serializer_class(self):
        if self.action == 'list':
            return ContratoListSerializer
        return ContratoDetailSerializer

    def _procesar_archivos_garantes(self, contrato, request):
        from django.core.files.storage import default_storage
        from django.core.files.base import ContentFile

        print('[DEBUG] request.FILES keys:', list(request.FILES.keys()))

        garantes = list(contrato.garantes or [])
        actualizado = False

        for i, garante in enumerate(garantes):
            archivo = request.FILES.get(f'garanteDocumentoArchivo_{i}')
            if archivo:
                import cloudinary.uploader
                extension = os.path.splitext(archivo.name)[1].lower().replace('.', '')
                nombre_sin_ext = os.path.splitext(archivo.name)[0]
                nombre_limpio = re.sub(r'[^a-zA-Z0-9_-]', '_', nombre_sin_ext)
                archivo.seek(0)
                result = cloudinary.uploader.upload(
                    archivo.read(),
                    folder="garantes",
                    public_id=f"{contrato.pk}_{i}_{nombre_limpio}",
                    format=extension,
                    resource_type="raw",
                    access_mode="public",
                    type="upload"
                )

                print("Cloudinary result:", result)
                print("URL guardada:", result['secure_url'])
                garante['documentoArchivo'] = result['secure_url']
                actualizado = True
                print(f'[DEBUG] garante {i} archivo guardado en: {garante["documentoArchivo"]}')
            elif garante.get('documentoArchivo') is None:
                # Preservar URL existente si ya tenía archivo (caso edición)
                pass

        if actualizado:
            contrato.garantes = garantes
            contrato.save(update_fields=['garantes'])

    def _procesar_contrato_pdf(self, contrato, request):
        """Procesar archivo PDF del contrato con Cloudinary API directa"""
        if 'contratoPdf' in request.FILES:
            archivo = request.FILES['contratoPdf']
            if archivo:
                import cloudinary.uploader
                extension = os.path.splitext(archivo.name)[1].lower().replace('.', '')
                nombre_sin_ext = os.path.splitext(archivo.name)[0]
                nombre_limpio = re.sub(r'[^a-zA-Z0-9_-]', '_', nombre_sin_ext)
                archivo.seek(0)
                result = cloudinary.uploader.upload(
                    archivo.read(),
                    folder="contratos/pdf",
                    public_id=f"{contrato.pk}_{nombre_limpio}",
                    format=extension,
                    resource_type="raw",
                    access_mode="public",
                    type="upload"
                )

                print("Cloudinary result (contrato PDF):", result)
                print("URL guardada (contrato PDF):", result['secure_url'])
                contrato.contratoPdf = result['secure_url']
                contrato.save(update_fields=['contratoPdf'])

    def perform_create(self, serializer):
        try:
            # Verificar que el usuario esté autenticado
            if not self.request.user or not self.request.user.is_authenticated:
                from rest_framework.exceptions import AuthenticationFailed
                raise AuthenticationFailed('Usuario no autenticado')
            
            contrato = serializer.save(usuario=self.request.user)
            self._procesar_archivos_garantes(contrato, self.request)
            self._procesar_contrato_pdf(contrato, self.request)
            creados = services.generar_meses(contrato)
            logger.info('Contrato %s creado con %d meses', contrato.pk, len(creados))
        except Exception as e:
            logger.error('ERROR al crear contrato: %s', str(e), exc_info=True)
            raise

    def perform_update(self, serializer):
        contrato = serializer.save()
        self._procesar_archivos_garantes(contrato, self.request)
        self._procesar_contrato_pdf(contrato, self.request)
        services.generar_meses(contrato, sobreescribir=False)

    def destroy(self, request, *args, **kwargs):
        """Soft delete — no borra físicamente el registro."""
        contrato = self.get_object()
        contrato.eliminado   = True
        contrato.eliminadoEn = timezone.now()
        contrato.save(update_fields=['eliminado', 'eliminadoEn'])
        logger.info('Contrato %s marcado como eliminado', contrato.pk)
        return Response(status=status.HTTP_204_NO_CONTENT)

    # ── POST /contratos/{id}/generar-meses/ ────────────────────────────────────
    @action(detail=True, methods=['post'], url_path='generar-meses')
    def generar_meses(self, request, pk=None):
        contrato     = self.get_object()
        sobreescribir = request.data.get('sobreescribir', False)
        creados      = services.generar_meses(contrato, sobreescribir=bool(sobreescribir))
        return Response({
            'message': f'{len(creados)} meses generados.',
            'creados': EstadoMensualSerializer(creados, many=True).data,
        }, status=status.HTTP_201_CREATED if creados else status.HTTP_200_OK)

    # ── GET /contratos/{id}/meses/ ─────────────────────────────────────────────
    @action(detail=True, methods=['get'], url_path='meses')
    def meses(self, request, pk=None):
        contrato = self.get_object()
        return Response(
            EstadoMensualSerializer(contrato.meses.all(), many=True).data
        )

    # ── PUT /contratos/{id}/meses/{mes}-{anio}/estado/ ─────────────────────────
    @action(detail=True, methods=['put'],
            url_path=r'meses/(?P<mes>\d+)-(?P<anio>\d+)/estado')
    def actualizar_estado_mes(self, request, pk=None, mes=None, anio=None):
        contrato = self.get_object()
        try:
            em = contrato.meses.get(mes=int(mes), anio=int(anio))
        except EstadoMensual.DoesNotExist:
            return Response(
                {'error': f'No existe el mes {mes}/{anio} para este contrato.'},
                status=status.HTTP_404_NOT_FOUND
            )
        serializer = EstadoMensualUpdateSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        em.estado = serializer.validated_data['estado']
        em.save(update_fields=['estado', 'updatedAt'])
        logger.info('Contrato %s mes %s/%s → %s', pk, mes, anio, em.estado)
        return Response(EstadoMensualSerializer(em).data)

    # ── POST /contratos/{id}/aplicar-aumento/ ─────────────────────────────────
    @action(detail=True, methods=['post'], url_path='aplicar-aumento')
    def aplicar_aumento(self, request, pk=None):
        """
        Calcula el aumento sugerido (no lo aplica aún).
        Para IPC/ICL consulta la API externa.
        """
        contrato   = self.get_object()
        serializer = AplicarAumentoSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data
        tipo = data['tipoAumento']

        indice_anterior = indice_nuevo = None

        if tipo in ('IPC', 'ICL', 'casa_propia'):
            resultado = obtener_indice(tipo)
            if resultado.get('error'):
                return Response({'error': resultado['error']}, status=status.HTTP_502_BAD_GATEWAY)
            porcentaje      = Decimal(str(resultado['valor']))
            indice_anterior = Decimal(str(resultado.get('anterior', 0)))
            indice_nuevo    = porcentaje
            monto_fijo      = None
        elif tipo == 'monto_fijo':
            monto_fijo = data.get('montoFijo')
            if monto_fijo is None:
                return Response({'error': 'montoFijo es requerido para tipo monto_fijo.'}, status=status.HTTP_400_BAD_REQUEST)
            porcentaje = None
            indice_anterior = None
            indice_nuevo = None
        else:
            porcentaje = data['porcentajeFijo']
            monto_fijo = None

        # Tomar monto base del primer mes pendiente desde mesDesde/anioDesde (o el primero disponible)
        mes_desde  = data.get('mesDesde')
        anio_desde = data.get('anioDesde')
        if mes_desde and anio_desde:
            mes_ref = contrato.meses.filter(mes=mes_desde - 1, anio=anio_desde).first()
        else:
            mes_ref = contrato.meses.order_by('anio', 'mes').first()

        monto_base = mes_ref.montoFinal if mes_ref else contrato.valorMensual

        if tipo == 'monto_fijo':
            monto_fijo = Decimal(str(monto_fijo))
            nuevo_monto = (monto_base + monto_fijo).quantize(Decimal('0.01'))
            porcentaje_sugerido = None
        else:
            nuevo_monto = services.calcular_nuevo_monto(monto_base, porcentaje)
            porcentaje_sugerido = porcentaje

        return Response({
            'tipoAumento':        tipo,
            'porcentajeSugerido': str(porcentaje_sugerido) if porcentaje_sugerido is not None else None,
            'montoBase':          str(monto_base),
            'nuevoMonto':         str(nuevo_monto),
            'indiceAnterior':     str(indice_anterior) if indice_anterior is not None else None,
            'indiceNuevo':        str(indice_nuevo)    if indice_nuevo    is not None else None,
            'montoFijo':          str(monto_fijo) if tipo == 'monto_fijo' else None,
        })

    # ── POST /contratos/{id}/confirmar-aumento/ ────────────────────────────────
    @action(detail=True, methods=['post'], url_path='confirmar-aumento')
    def confirmar_aumento(self, request, pk=None):
        """Aplica el aumento confirmado a los meses correspondientes."""
        print("DATA recibida:", request.data)
        print("tipoAumento:", request.data.get('tipoAumento'))
        print("porcentajeAumento:", request.data.get('porcentajeAumento'))
        contrato   = self.get_object()
        serializer = ConfirmarAumentoSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data

        resultados = services.aplicar_aumento(
            contrato        = contrato,
            tipo_aumento    = data['tipoAumento'],
            porcentaje      = data['porcentajeAumento'],
            monto_fijo      = data.get('montoFijo'),
            indice_anterior = data.get('indiceAnterior'),
            indice_nuevo    = data.get('indiceNuevo'),
            mes_desde       = data.get('mesDesde'),
            anio_desde      = data.get('anioDesde'),
            razon           = data.get('razon', ''),
            aplicado_por    = data.get('aplicadoPor', ''),
        )
        return Response({
            'message':    f'Aumento aplicado a {len(resultados)} meses.',
            'resultados': resultados,
        })

    # ── POST /contratos/{id}/aplicar-aumento-mora/ ─────────────────────────────
    @action(detail=True, methods=['post'], url_path='aplicar-aumento-mora')
    def aplicar_mora(self, request, pk=None):
        contrato   = self.get_object()
        serializer = AplicarMoraSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data

        if not contrato.valorInteresMora:
            return Response(
                {'error': 'Este contrato no tiene configurado un interés de mora.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            resultados = services.aplicar_mora(
                contrato     = contrato,
                aplicado_por = data.get('aplicadoPor', ''),
                mes          = data.get('mes'),
                anio         = data.get('anio'),
                dias_atraso  = data.get('diasAtraso'),
                recargo_mora = data.get('recargoMora'),
            )
        except EstadoMensual.DoesNotExist:
            return Response(
                {'error': 'No existe el mes indicado para este contrato.'},
                status=status.HTTP_404_NOT_FOUND
            )
        except ValueError as exc:
            return Response({'error': str(exc)}, status=status.HTTP_409_CONFLICT)

        return Response({
            'message':    f'Mora aplicada a {len(resultados)} meses.',
            'resultados': resultados,
        })

    # ── GET /contratos/{id}/resumen-financiero/ ────────────────────────────────
    @action(detail=True, methods=['get'], url_path='resumen-financiero')
    def resumen_financiero(self, request, pk=None):
        contrato = self.get_object()
        return Response(services.resumen_financiero(contrato))

    # ── GET /contratos/buscar/ ────────────────────────────────────────────────
    @action(detail=False, methods=['get'], url_path='buscar')
    def buscar(self, request):
        """Búsqueda avanzada — usa los mismos filtros que el listado."""
        qs         = self.filter_queryset(self.get_queryset())
        serializer = ContratoListSerializer(qs, many=True)
        return Response(serializer.data)
    
    # ── POST /contratos/{id}/recalcular-montos/ ────────────────────────────────
    @action(detail=True, methods=['post'], url_path='recalcular-montos')
    def recalcular_montos(self, request, pk=None):
        contrato = self.get_object()

        nuevo_valor = Decimal(str(contrato.valorMensual))

        actualizados = 0
        for mes in contrato.meses.order_by('anio', 'mes'):
            porcentaje_acumulado = Decimal('1')
            for aumento in mes.aumentos.exclude(tipoAumento='mora'):
                porcentaje_acumulado *= (1 + aumento.porcentajeAumento / 100)

            monto_sin_mora = (nuevo_valor * porcentaje_acumulado).quantize(Decimal('0.01'))
            mes.montoBase  = nuevo_valor
            mes.montoFinal = (
                (monto_sin_mora + mes.recargo_mora).quantize(Decimal('0.01'))
                if mes.mora_aplicada and mes.recargo_mora
                else monto_sin_mora
            )
            mes.save(update_fields=['montoBase', 'montoFinal', 'updatedAt'])
            actualizados += 1

        logger.info('Contrato %s — %d meses recalculados con base %s', pk, actualizados, nuevo_valor)
        return Response({'ok': True, 'meses_actualizados': actualizados})

    # ── POST /contratos/{id}/recibo/ ───────────────────────────────────────────
    @action(detail=True, methods=['post'], url_path='recibo')
    def generar_recibo(self, request, pk=None):
        """Genera un recibo de pago en formato .docx"""
        contrato = self.get_object()
        serializer = ReciboSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data

        monto_alquiler  = Decimal(str(data['montoAlquiler']))
        conceptos       = list(contrato.conceptosExtras or []) or data.get('conceptosExtras') or []
        if conceptos:
            total_extras = sum(Decimal(str(item.get('precio', item.get('valor', 0)))) for item in conceptos)
        else:
            total_extras = Decimal(str(data.get('totalExtras') or 0))
        recargo_mora   = Decimal(str(data.get('recargoMora') or 0))
        dias_atraso    = data.get('diasAtraso') or 0
        valor_interes  = Decimal(str(data.get('valorInteresMora') or 0))
        total_monto    = (monto_alquiler + total_extras + recargo_mora).quantize(Decimal('0.01'))

        doc = Document()

        import os
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'logo-inmobiliaria-recibo.jpg')
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(4.0))

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run("Martires Riocuartenses N° 1395 – X5800 – Rio Cuarto – Córdoba.\n")
        header.add_run("9 de Julio Nº 483-x6125-Serrano-Córdoba.\n")
        header.add_run("Tel: 358 4864404 o 3385 465877 - E-Mail: inmobiliariagiordanoconti@gmail.com")

        doc.add_paragraph()

        def formatear_monto(monto):
            return f"{int(monto):,}".replace(',', '.') + f",{int((monto % 1) * 100):02d}"

        def linea(label, valor=''):
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            pPr = p._p.get_or_add_pPr()
            tabs_el = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '9350')
            tab.set(qn('w:leader'), 'dot')
            tabs_el.append(tab)
            pPr.append(tabs_el)
            p.add_run(label)
            p.add_run('\t')
            p.add_run(valor)
            return p

        meses_espanol = {
            1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL', 5: 'MAYO', 6: 'JUNIO',
            7: 'JULIO', 8: 'AGOSTO', 9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
        }
        fecha_formateada = (
            f"{contrato.fechaInicio.day} DE "
            f"{meses_espanol[contrato.fechaInicio.month]} {contrato.fechaInicio.year}"
        )

        direccion_completa = contrato.direccion
        if contrato.piso and contrato.piso.strip() not in ('', '-'):
            direccion_completa += f" Piso {contrato.piso}"
        if contrato.departamento and contrato.departamento.strip() not in ('', '-'):
            direccion_completa += f" Dpto. {contrato.departamento}"

        monto_en_letras = convertir_monto_a_letras(monto_alquiler)
        texto_principal = (
            f"Recibo del Sr./Sra. {contrato.inquilinoNombre.upper()}, DNI Nº {contrato.inquilinoDni}, "
            f"TEL Nº {contrato.inquilinoTelefono}, EMAIL {getattr(contrato, 'inquilinoEmail', '')}, "
            f"de la ciudad de {contrato.localidad}, provincia de {contrato.provincia} "
            f"la suma de pesos: {monto_en_letras} ($ {formatear_monto(monto_alquiler)}), "
            f"por cuenta y orden de terceros, conforme contrato de locación con fecha {fecha_formateada}, "
            f"con relación al inmueble ubicado en {direccion_completa}, en concepto de:"
        )
        doc.add_paragraph(texto_principal)
        doc.add_paragraph()

        linea(f"-ALQUILER {data['mes'].upper()} {data['anio']}", f"$ {formatear_monto(monto_alquiler)}.")
        linea("-EMOS", "Abona locataria.")
        linea("-MUNICIPAL", "Abona locataria.")
        if conceptos:
            for item in conceptos:
                nombre = str(item.get('nombre', item.get('concepto', 'EXTRA'))).upper()
                valor  = Decimal(str(item.get('precio', item.get('valor', 0))))
                if valor > 0:
                    linea(f"-{nombre}", f"$ {formatear_monto(valor)}.")
                else:
                    linea(f"-{nombre}", "Abona la locataria.")
        elif total_extras > 0:
            linea("-EXPENSAS", f"$ {formatear_monto(total_extras)}.")
        else:
            linea("-EXPENSAS", "Abona la locataria.")
        if recargo_mora > 0:
            linea(f"-MORA ({dias_atraso} días x {valor_interes}%)", f"$ {formatear_monto(recargo_mora)}.")
        linea("SUBTOTAL", f"$ {formatear_monto(total_monto)}.")
        linea("-DESCUENTO", "")
        linea("TOTAL", f"$ {formatear_monto(total_monto)}.")

        doc.add_paragraph()

        firma = doc.add_paragraph()
        firma_run = firma.add_run("Recibí Conforme: PAGO RECIBIDO MEDIANTE TRANSFERENCIA BANCARIA")
        firma_run.bold = True
        
        # Guardar documento en memoria
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Preparar respuesta
        # Limpiar nombre del inquilino para usarlo en el nombre de archivo
        nombre_inquilino_limpio = contrato.inquilinoNombre.replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"recibo_{nombre_inquilino_limpio}_{data['mes']}_{data['anio']}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info('Recibo generado para contrato %s - %s %s', pk, data['mes'], data['anio'])
        return response

    # ── POST /contratos/{id}/recibo-propietario/ ──────────────────────────────
    @action(detail=True, methods=['post'], url_path='recibo-propietario')
    def generar_recibo_propietario(self, request, pk=None):
        """Genera un recibo de honorarios para el propietario en formato .docx"""
        contrato = self.get_object()
        serializer = ReciboSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data

        monto_alquiler   = Decimal(str(data['montoAlquiler']))
        honorarios_pct   = Decimal(str(contrato.honorarios or 0))
        monto_honorarios = (monto_alquiler * honorarios_pct / 100).quantize(Decimal('0.01'))
        subtotal         = monto_alquiler.quantize(Decimal('0.01'))
        total_propietario = (subtotal - monto_honorarios).quantize(Decimal('0.01'))

        doc = Document()

        import os
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'logo-inmobiliaria-recibo.jpg')
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(4.0))

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run("Martires Riocuartenses N° 1395 – X5800 – Rio Cuarto – Córdoba.\n")
        header.add_run("9 de Julio Nº 483-x6125-Serrano-Córdoba.\n")
        header.add_run("Tel: 358 4864404 o 3385 465877 - E-Mail: inmobiliariagiordanoconti@gmail.com")

        doc.add_paragraph()

        def formatear_monto(monto):
            return f"{int(monto):,}".replace(',', '.') + f",{int((monto % 1) * 100):02d}"

        def linea(label, valor=''):
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            pPr = p._p.get_or_add_pPr()
            tabs_el = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '9350')
            tab.set(qn('w:leader'), 'dot')
            tabs_el.append(tab)
            pPr.append(tabs_el)
            p.add_run(label)
            p.add_run('\t')
            p.add_run(valor)
            return p

        meses_espanol = {
            1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL', 5: 'MAYO', 6: 'JUNIO',
            7: 'JULIO', 8: 'AGOSTO', 9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
        }
        fecha_formateada = (
            f"{contrato.fechaInicio.day} DE "
            f"{meses_espanol[contrato.fechaInicio.month]} {contrato.fechaInicio.year}"
        )

        direccion_completa = contrato.direccion
        if contrato.piso and contrato.piso.strip() not in ('', '-'):
            direccion_completa += f" Piso {contrato.piso}"
        if contrato.departamento and contrato.departamento.strip() not in ('', '-'):
            direccion_completa += f" Dpto. {contrato.departamento}"

        monto_en_letras = convertir_monto_a_letras(monto_alquiler)
        texto_principal = (
            f"Recibo del Sr./Sra. {contrato.inquilinoNombre.upper()}, DNI Nº {contrato.inquilinoDni}, "
            f"TEL Nº {contrato.inquilinoTelefono}, EMAIL {getattr(contrato, 'inquilinoEmail', '')}, "
            f"de la ciudad de {contrato.localidad}, provincia de {contrato.provincia} "
            f"la suma de pesos: {monto_en_letras} ($ {formatear_monto(monto_alquiler)}), "
            f"por cuenta y orden de terceros, conforme contrato de locación con fecha {fecha_formateada}, "
            f"con relación al inmueble ubicado en {direccion_completa}, en concepto de:"
        )
        doc.add_paragraph(texto_principal)
        doc.add_paragraph()

        linea(f"-ALQUILER {data['mes'].upper()} {data['anio']}", f"$ {formatear_monto(monto_alquiler)}.")
        linea("-EMOS", "Abona locataria.")
        linea("-MUNICIPAL", "Abona locataria.")
        linea("SUBTOTAL", f"$ {formatear_monto(subtotal)}.")
        linea(f"-GTOS ADMINIST. {honorarios_pct}%", f"$ {formatear_monto(monto_honorarios)}.")
        linea("TOTAL", f"$ {formatear_monto(total_propietario)}.")

        doc.add_paragraph()

        firma = doc.add_paragraph()
        firma.add_run("Recibí Conforme: PAGO REALIZADO MEDIANTE TRANSFERENCIA BANCARIA").bold = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"recibo_propietario_{data['mes']}_{data['anio']}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        logger.info('Recibo propietario generado para contrato %s - %s %s', pk, data['mes'], data['anio'])
        return response

    @action(detail=True, methods=['get'], url_path='documento-garante/(?P<indice>[0-9]+)')
    def documento_garante(self, request, pk=None, indice=None):
        contrato = self.get_object()
        garantes = contrato.garantes or []
        try:
            garante = garantes[int(indice)]
        except IndexError:
            return Response({'error': 'Garante no encontrado'}, status=status.HTTP_404_NOT_FOUND)

        url = garante.get('documentoArchivo')
        if not url:
            return Response({'error': 'El garante no tiene documento'}, status=status.HTTP_404_NOT_FOUND)

        r = req.get(url, timeout=30)
        content_type = r.headers.get('Content-Type', 'application/octet-stream')

        response = HttpResponse(r.content, content_type=content_type)
        response['Content-Disposition'] = 'inline; filename="documento"'
        return response
